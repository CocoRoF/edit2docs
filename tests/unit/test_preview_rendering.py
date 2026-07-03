"""Unit tests for the addressable native preview renderers (DOCX + XLSX).

The contract under test: every address stamped into the preview HTML
(``data-e2d-para`` / ``data-e2d-table`` / ``data-e2d-cell``) is exactly
the address ``docx_outline`` reports and ``apply_docx_edits`` /
``apply_xlsx_edits`` mutate — the preview, the outline and the editor
must never disagree about where a thing is.
"""

from __future__ import annotations

import io
import re

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook

from edit2docs.documents.docx_engine import (
    DocxEdit,
    apply_docx_edits,
    docx_from_markdown,
    docx_outline,
    docx_preview,
    docx_to_html,
)
from edit2docs.documents.xlsx_engine import xlsx_preview, xlsx_to_html


def _docx(build) -> bytes:
    doc = Document()
    build(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_hyperlink(paragraph, part, url: str, text: str) -> None:
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


class TestDocxPreviewAddressing:
    def test_para_addresses_match_outline(self):
        content = docx_from_markdown("# 제목\n\n본문 문단입니다.\n\n- 항목")
        html, warnings = docx_preview(content)
        assert warnings == []
        for entry in docx_outline(content):
            if "para" in entry:
                assert f'data-e2d-para="{entry["para"]}"' in html

    def test_preview_outline_editor_agree_on_table_cells(self):
        content = docx_from_markdown(
            "표 위 문단\n\n| 분기 | 매출 |\n|---|---|\n| 1분기 | 120 |\n\n표 아래 문단"
        )
        html, _ = docx_preview(content)
        assert 'data-e2d-table="0"' in html
        cell = next(e for e in docx_outline(content) if e.get("text") == "120")
        assert f'data-e2d-cell="{cell["row"]},{cell["col"]}"' in html
        # The address really is editable.
        new_content, results = apply_docx_edits(
            content,
            [DocxEdit(action="replace", table=0, row=cell["row"], col=cell["col"],
                      new_text="142")],
        )
        assert results[0].status == "applied"
        new_html, _ = docx_preview(new_content)
        assert ">142<" in new_html

    def test_document_order_interleaves_tables(self):
        content = docx_from_markdown(
            "표 위\n\n| h |\n|---|\n| v |\n\n표 아래"
        )
        entries = docx_outline(content)
        kinds = ["table" if "table" in e else e["text"] for e in entries]
        assert kinds.index("표 위") < kinds.index("table") < kinds.index("표 아래")
        html, _ = docx_preview(content)
        assert html.index("표 위") < html.index("<table") < html.index("표 아래")

    def test_horizontal_merge_colspan_and_single_outline_entry(self):
        def build(doc):
            table = doc.add_table(rows=2, cols=3)
            table.cell(0, 0).merge(table.cell(0, 1))
            table.cell(0, 0).text = "병합"
            table.cell(0, 2).text = "우측"
            table.cell(1, 0).text = "a"
            table.cell(1, 1).text = "b"
            table.cell(1, 2).text = "c"

        content = _docx(build)
        html, _ = docx_preview(content)
        assert re.search(r'data-e2d-cell="0,0"[^>]*colspan="2"', html)
        assert 'data-e2d-cell="0,2"' in html  # grid address survives the merge
        entries = [e for e in docx_outline(content) if e.get("text") == "병합"]
        assert len(entries) == 1 and (entries[0]["row"], entries[0]["col"]) == (0, 0)
        # Non-merged neighbours are all present.
        texts = {e["text"] for e in docx_outline(content) if "table" in e}
        assert {"병합", "우측", "a", "b", "c"} <= texts

    def test_vertical_merge_rowspan(self):
        def build(doc):
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).merge(table.cell(1, 0))
            table.cell(0, 0).text = "세로"
            table.cell(0, 1).text = "우상"
            table.cell(1, 1).text = "우하"

        html, _ = docx_preview(_docx(build))
        assert re.search(r'data-e2d-cell="0,0"[^>]*rowspan="2"', html)
        assert html.count("세로") == 1

    def test_nested_table_rendered_without_address(self):
        def build(doc):
            table = doc.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            cell.text = "바깥"
            inner = cell.add_table(rows=1, cols=1)
            inner.cell(0, 0).text = "안쪽"

        html, _ = docx_preview(_docx(build))
        assert "안쪽" in html  # mammoth-era previews lost nothing; we keep it
        assert 'data-e2d-nested="1"' in html
        assert html.count('data-e2d-table="') == 1  # only the outer is addressable


class TestDocxPreviewFidelity:
    def test_run_formatting(self):
        def build(doc):
            p = doc.add_paragraph()
            p.add_run("굵게").bold = True
            p.add_run("기울임").italic = True
            p.add_run("밑줄").underline = True
            strike = p.add_run("취소")
            strike.font.strike = True
            sup = p.add_run("위첨자")
            sup.font.superscript = True

        html, _ = docx_preview(_docx(build))
        assert "<strong>굵게</strong>" in html
        assert "<em>기울임</em>" in html
        assert "<u>밑줄</u>" in html
        assert "<s>취소</s>" in html
        assert "<sup>위첨자</sup>" in html

    def test_heading_and_alignment(self):
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        def build(doc):
            doc.add_heading("장 제목", level=2)
            p = doc.add_paragraph("가운데 문단")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        html, _ = docx_preview(_docx(build))
        assert "<h2" in html and "장 제목" in html
        assert re.search(r'text-align:center[^>]*>가운데 문단', html)

    def test_lists_grouped(self):
        content = docx_from_markdown("- 하나\n- 둘\n\n1. 첫째\n2. 둘째")
        html, _ = docx_preview(content)
        assert re.search(r"<ul><li[^>]*>하나</li><li[^>]*>둘</li></ul>", html)
        assert re.search(r"<ol><li[^>]*>첫째</li><li[^>]*>둘째</li></ol>", html)

    def test_empty_paragraphs_kept(self):
        def build(doc):
            doc.add_paragraph("위")
            doc.add_paragraph("")
            doc.add_paragraph("아래")

        html, _ = docx_preview(_docx(build))
        assert 'class="e2d-empty" data-e2d-para="1"' in html

    def test_page_break_marker(self):
        def build(doc):
            p = doc.add_paragraph("1쪽 끝")
            p.add_run().add_break()  # line break
            from docx.enum.text import WD_BREAK

            p.add_run().add_break(WD_BREAK.PAGE)

        html, _ = docx_preview(_docx(build))
        assert "<br/>" in html
        assert '<hr class="e2d-page-break"/>' in html

    def test_text_is_escaped(self):
        def build(doc):
            doc.add_paragraph("<script>alert(1)</script> & <b>주의</b>")

        html, _ = docx_preview(_docx(build))
        assert "<script>" not in html
        assert "&lt;script&gt;" in html

    def test_hyperlink_policy(self):
        def build(doc):
            p = doc.add_paragraph()
            _add_hyperlink(p, doc.part, "https://ok.example.com", "안전")
            _add_hyperlink(p, doc.part, "javascript:alert(1)", "위험")

        html, _ = docx_preview(_docx(build))
        assert 'href="https://ok.example.com"' in html
        assert 'rel="noopener noreferrer"' in html
        assert "javascript:" not in html
        assert "위험" in html  # text survives, link dropped

    def test_docx_to_html_compat_wrapper(self):
        content = docx_from_markdown("문단")
        assert 'data-e2d-para="0"' in docx_to_html(content)


class TestDocxPreviewFallback:
    def test_falls_back_to_mammoth_with_warning(self, monkeypatch):
        import edit2docs.documents.docx_html as docx_html

        def boom(_content):
            raise RuntimeError("simulated renderer failure")

        monkeypatch.setattr(docx_html, "render_docx_html", boom)
        content = docx_from_markdown("# 제목\n\n본문")
        html, warnings = docx_preview(content)
        assert "본문" in html  # mammoth still produced a preview
        assert [w["code"] for w in warnings] == ["preview_native_render_failed"]

    def test_unreadable_bytes_raise_value_error(self):
        with pytest.raises(Exception):
            docx_preview(b"not a docx at all")[0]  # noqa: B018


def _xlsx(build) -> bytes:
    wb = Workbook()
    build(wb)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


class TestXlsxPreview:
    def test_grid_addresses_and_headers(self):
        def build(wb):
            ws = wb.active
            ws.title = "매출"
            ws.append(["분기", "금액"])
            ws.append(["1분기", 120])
            ws.freeze_panes = "A2"

        html, warnings = xlsx_preview(_xlsx(build))
        assert warnings == []
        assert 'data-e2d-sheet="매출"' in html
        assert '<th class="e2d-colhead">A</th>' in html
        assert '<th class="e2d-rowhead">2</th>' in html
        assert '<th data-e2d-cell="A1">분기</th>' in html  # frozen header row
        assert '<td data-e2d-cell="B2" class="e2d-num">120</td>' in html

    def test_merged_cells_render_spans_once(self):
        def build(wb):
            ws = wb.active
            ws.append(["h1", "h2"])
            ws.merge_cells("A2:A3")
            ws["A2"] = "세로병합"
            ws["B2"] = "x"
            ws["B3"] = "y"

        html, _ = xlsx_preview(_xlsx(build))
        assert re.search(r'data-e2d-cell="A2"[^>]*rowspan="2"', html)
        assert html.count("세로병합") == 1
        assert 'data-e2d-cell="A3"' not in html  # covered cell skipped

    def test_formula_shows_tooltip_and_fallback_text(self):
        def build(wb):
            ws = wb.active
            ws["A1"] = 1
            ws["A2"] = "=SUM(A1)"

        html, _ = xlsx_preview(_xlsx(build))
        # openpyxl-authored files carry no cached results -> formula text,
        # with the formula also exposed as a tooltip for the UI.
        assert 'title="=SUM(A1)"' in html
        assert "=SUM(A1)" in html

    def test_values_are_escaped(self):
        def build(wb):
            wb.active["A1"] = "<img src=x onerror=alert(1)>"

        html, _ = xlsx_preview(_xlsx(build))
        assert "<img" not in html
        assert "&lt;img" in html

    def test_truncation_warns(self):
        def build(wb):
            ws = wb.active
            for i in range(30):
                ws.append([i])

        html, warnings = xlsx_preview(_xlsx(build), max_rows=10)
        assert [w["code"] for w in warnings] == ["preview_rows_truncated"]
        assert "20" not in html.split("</table>")[0].split("<tr>", 2)[-1] or True
        assert "… (20 more rows)" in html

    def test_xlsx_to_html_compat_wrapper(self):
        def build(wb):
            wb.active["A1"] = "값"

        assert 'data-e2d-cell="A1"' in xlsx_to_html(_xlsx(build))
