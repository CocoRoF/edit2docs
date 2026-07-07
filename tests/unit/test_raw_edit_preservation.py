"""Regression tests: the raw-layer edit paths must NOT destroy data.

These port the two audited P0 data-destruction scenarios that the old
openpyxl / python-docx round-trips triggered on every edit:

* **P0-2 (xlsx)** — ``openpyxl load→save`` rewrote the whole package,
  silently dropping sparkline ``extLst`` blocks, custom XML, chart style
  parts and cached formula values. The raw ``set_cell`` path rewrites
  only the touched worksheet, so all of those survive byte-identical.
* **P0-3 (docx)** — cell/paragraph replace rebuilt runs, destroying
  in-cell images, collapsing multi-paragraph cells, dropping run
  formatting and deleting hyperlink elements wholesale. The raw
  replace path is run-preserving.
"""

from __future__ import annotations

import io
import zipfile

import openpyxl
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from lxml import etree
from PIL import Image as PilImage

from edit2docs.documents.docx_engine import (
    DocxEdit,
    apply_docx_edits,
    docx_outline,
)
from edit2docs.documents.xlsx_engine import (
    XlsxEdit,
    apply_xlsx_edits,
    xlsx_preview,
)

# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------


def _parts(data: bytes) -> dict[str, bytes]:
    z = zipfile.ZipFile(io.BytesIO(data))
    return {n: z.read(n) for n in z.namelist()}


def _diff_parts(a: bytes, b: bytes) -> set[str]:
    pa, pb = _parts(a), _parts(b)
    names = set(pa) | set(pb)
    return {n for n in names if pa.get(n) != pb.get(n)}


def _rezip(data: bytes, *, replace=None, add=None) -> bytes:
    z = zipfile.ZipFile(io.BytesIO(data))
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as znew:
        for n in z.namelist():
            znew.writestr(n, (replace or {}).get(n, z.read(n)))
        for n, content in (add or {}).items():
            znew.writestr(n, content)
    return out.getvalue()


def _png_bytes() -> bytes:
    buf = io.BytesIO()
    PilImage.new("RGB", (4, 4), (30, 100, 218)).save(buf, format="PNG")
    return buf.getvalue()


# ---------------------------------------------------------------------------
# xlsx fixtures
# ---------------------------------------------------------------------------

CONTENT_TYPES = "[Content_Types].xml"
WORKBOOK = "xl/workbook.xml"
WORKBOOK_RELS = "xl/_rels/workbook.xml.rels"


def _build_workbook() -> bytes:
    """Two sheets: Sales (with a formula) + Notes."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sales"
    ws["A1"], ws["B1"] = "Region", "Q1"
    ws["A2"], ws["B2"] = "North", 300
    ws["A3"], ws["B3"] = "South", 275
    ws["B4"] = "=SUM(B2:B3)"
    wb.create_sheet("Notes")["A1"] = "메모"
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _sheet_part(data: bytes, sheet_name: str) -> str:
    """Worksheet part name for *sheet_name* (via workbook rels)."""
    parts = _parts(data)
    wb_root = etree.fromstring(parts[WORKBOOK])
    ns_s = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
    ns_r = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    rid = None
    for sheet in wb_root.findall(f"{{{ns_s}}}sheets/{{{ns_s}}}sheet"):
        if sheet.get("name") == sheet_name:
            rid = sheet.get(f"{{{ns_r}}}id")
    rels = etree.fromstring(parts[WORKBOOK_RELS])
    ns_p = "http://schemas.openxmlformats.org/package/2006/relationships"
    for rel in rels.findall(f"{{{ns_p}}}Relationship"):
        if rel.get("Id") == rid:
            return "xl/" + rel.get("Target").lstrip("/").removeprefix("xl/")
    raise AssertionError(f"no part for sheet {sheet_name!r}")


_SPARKLINE_EXT = (
    b'<extLst><ext xmlns:x14="http://schemas.microsoft.com/office/'
    b'spreadsheetml/2009/9/main" uri="{05C60535-1F16-4fd2-B633-F4F36F0B64E0}">'
    b'<x14:sparklineGroups xmlns:xm="http://schemas.microsoft.com/office/'
    b'excel/2006/main"><x14:sparklineGroup displayEmptyCellsAs="gap">'
    b'<x14:colorSeries rgb="FF376092"/><x14:sparklines><x14:sparkline>'
    b"<xm:f>Sales!B2:B3</xm:f><xm:sqref>A5</xm:sqref></x14:sparkline>"
    b"</x14:sparklines></x14:sparklineGroup></x14:sparklineGroups>"
    b"</ext></extLst>"
)

FOREIGN_PARTS = [
    "xl/charts/style1.xml",
    "xl/charts/colors1.xml",
    "customXml/item1.xml",
    "customXml/itemProps1.xml",
    "customXml/_rels/item1.xml.rels",
]


def _inject_foreign_features(base: bytes) -> bytes:
    """Add the parts openpyxl round-trips destroy: chart style/colors,
    custom XML (with rels + content types), and a sparkline extLst inside
    the Notes worksheet."""
    parts = _parts(base)
    notes_part = _sheet_part(base, "Notes")
    ct = parts[CONTENT_TYPES].replace(
        b"</Types>",
        b'<Override PartName="/xl/charts/style1.xml" ContentType="application/'
        b'vnd.ms-office.chartstyle+xml"/>'
        b'<Override PartName="/xl/charts/colors1.xml" ContentType="application/'
        b'vnd.ms-office.chartcolorstyle+xml"/>'
        b'<Override PartName="/customXml/itemProps1.xml" ContentType="application/'
        b'vnd.openxmlformats-officedocument.customXmlProperties+xml"/>'
        b"</Types>",
    )
    wb_rels = parts[WORKBOOK_RELS].replace(
        b"</Relationships>",
        b'<Relationship Id="rId98" Type="http://schemas.openxmlformats.org/'
        b'officeDocument/2006/relationships/customXml" '
        b'Target="../customXml/item1.xml"/></Relationships>',
    )
    notes = parts[notes_part].replace(
        b"</worksheet>", _SPARKLINE_EXT + b"</worksheet>"
    )
    item_rels = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        b'<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/'
        b'relationships">'
        b'<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/'
        b'officeDocument/2006/relationships/customXmlProps" '
        b'Target="itemProps1.xml"/></Relationships>'
    )
    return _rezip(
        base,
        replace={CONTENT_TYPES: ct, WORKBOOK_RELS: wb_rels, notes_part: notes},
        add={
            "xl/charts/style1.xml": (
                b'<cs:chartStyle xmlns:cs="http://schemas.microsoft.com/office/'
                b'drawing/2012/chartStyle" id="201"/>'
            ),
            "xl/charts/colors1.xml": (
                b'<cs:colorStyle xmlns:cs="http://schemas.microsoft.com/office/'
                b'drawing/2012/chartStyle" meth="cycle" id="10"/>'
            ),
            "customXml/item1.xml": (
                "<meta xmlns='urn:example:meta'><owner>홍길동</owner></meta>"
            ).encode(),
            "customXml/itemProps1.xml": (
                b'<ds:datastoreItem xmlns:ds="http://schemas.openxmlformats.org/'
                b'officeDocument/2006/customXml" ds:itemID="{A1B2C3D4-0000-0000-'
                b'0000-000000000001}"/>'
            ),
            "customXml/_rels/item1.xml.rels": item_rels,
        },
    )


def _set_cached_formula(base: bytes, value: str) -> bytes:
    """Give the Sales!B4 formula cell a cached ``<v>`` (Excel writes one;
    openpyxl leaves ``<v></v>``)."""
    part = _sheet_part(base, "Sales")
    parts = _parts(base)
    s = parts[part]
    if b"<f>SUM(B2:B3)</f><v></v>" in s:
        s = s.replace(
            b"<f>SUM(B2:B3)</f><v></v>",
            b"<f>SUM(B2:B3)</f><v>" + value.encode() + b"</v>",
        )
    else:  # openpyxl variant without the empty <v>
        s = s.replace(
            b"<f>SUM(B2:B3)</f>",
            b"<f>SUM(B2:B3)</f><v>" + value.encode() + b"</v>",
        )
    return _rezip(base, replace={part: s})


# ---------------------------------------------------------------------------
# xlsx tests (P0-2)
# ---------------------------------------------------------------------------


class TestXlsxPreservation:
    def test_injected_parts_survive_set_cell_byte_identical(self):
        """P0-2 proof: sparkline extLst + customXml + chart style parts
        survive one set_cell BYTE-IDENTICAL."""
        injected = _inject_foreign_features(_build_workbook())
        new, results = apply_xlsx_edits(
            injected,
            [XlsxEdit(action="set_cell", sheet="Sales", cell="B2", value=999)],
        )
        assert results[0].status == "applied"

        before, after = _parts(injected), _parts(new)
        # Every foreign part is present and byte-identical.
        for name in FOREIGN_PARTS:
            assert name in after, f"{name} dropped"
            assert after[name] == before[name], f"{name} was rewritten"
        # The sparkline extLst rides on the Notes sheet (untouched).
        notes_part = _sheet_part(injected, "Notes")
        assert after[notes_part] == before[notes_part]
        assert b"sparklineGroups" in after[notes_part]

    def test_set_cell_touches_only_its_sheet(self):
        """Byte-preservation: a single set_cell in a formula-free workbook
        rewrites ONLY that worksheet part."""
        wb = openpyxl.Workbook()
        wb.active.title = "Data"
        wb.active["A1"] = 1
        wb.create_sheet("Other")["A1"] = "x"
        buf = io.BytesIO()
        wb.save(buf)
        base = buf.getvalue()

        new, results = apply_xlsx_edits(
            base, [XlsxEdit(action="set_cell", sheet="Data", cell="A1", value=42)]
        )
        assert results[0].status == "applied"
        assert _diff_parts(base, new) == {_sheet_part(base, "Data")}

    def test_set_cell_with_formula_touches_only_sheet_and_workbook(self):
        """With a formula present, set_cell also flips calcPr in
        workbook.xml — but nothing else."""
        base = _set_cached_formula(_build_workbook(), "575")
        new, results = apply_xlsx_edits(
            base, [XlsxEdit(action="set_cell", sheet="Sales", cell="A1", value="지역")]
        )
        assert results[0].status == "applied"
        assert _diff_parts(base, new) <= {_sheet_part(base, "Sales"), WORKBOOK}

    def test_cached_formula_value_survives_unrelated_edit(self):
        """Preview shows the cached 575, not the raw =SUM, after editing an
        unrelated cell (the old openpyxl round-trip blanked the cache)."""
        base = _set_cached_formula(_build_workbook(), "575")
        new, results = apply_xlsx_edits(
            base, [XlsxEdit(action="set_cell", sheet="Sales", cell="A1", value="지역")]
        )
        assert results[0].status == "applied"

        html, _ = xlsx_preview(new)
        assert ">575<" in html
        # The cell BODY shows the cached value, not the raw formula (the
        # formula is allowed only in the title tooltip: `title="=SUM..."`).
        assert ">=SUM" not in html

        # And the cached value is really still in the stored bytes.
        wb_values = openpyxl.load_workbook(io.BytesIO(new), data_only=True)
        assert wb_values["Sales"]["B4"].value == 575

    def test_style_preserved_on_set_cell(self):
        """The edited cell keeps its style index (@s) — proof the raw path
        writes surgically rather than reconstructing the cell."""
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "S"
        ws["A1"] = "old"
        ws["A1"].number_format = "#,##0.00"
        buf = io.BytesIO()
        wb.save(buf)
        base = buf.getvalue()

        new, results = apply_xlsx_edits(
            base, [XlsxEdit(action="set_cell", sheet="S", cell="A1", value="new")]
        )
        assert results[0].status == "applied"
        reopened = openpyxl.load_workbook(io.BytesIO(new))
        assert reopened["S"]["A1"].value == "new"
        assert reopened["S"]["A1"].number_format == "#,##0.00"


# ---------------------------------------------------------------------------
# docx fixtures
# ---------------------------------------------------------------------------


def _doc_bytes(document: Document) -> bytes:
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _document_xml(data: bytes) -> etree._Element:
    return etree.fromstring(_parts(data)["word/document.xml"])


def _media_parts(data: bytes) -> set[str]:
    return {n for n in _parts(data) if n.startswith("word/media/")}


# ---------------------------------------------------------------------------
# docx tests (P0-3)
# ---------------------------------------------------------------------------


class TestDocxPreservation:
    def test_cell_replace_preserves_image_and_layout(self):
        """Cell = [text, image-paragraph, text]. A cell replace must keep
        the in-cell image (drawing + media) and not collapse the layout."""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.paragraphs[0].add_run("위 텍스트")
        img_p = cell.add_paragraph()
        img_p.add_run().add_picture(io.BytesIO(_png_bytes()), width=Inches(0.5))
        cell.add_paragraph().add_run("아래 텍스트")
        base = _doc_bytes(doc)
        assert len(_media_parts(base)) == 1

        new, results = apply_docx_edits(
            base,
            [DocxEdit(action="replace", table=0, row=0, col=0, new_text="새 텍스트")],
        )
        assert results[0].status == "applied"

        # Media survived; the drawing element is still in the cell.
        assert _media_parts(new) == _media_parts(base)
        root = _document_xml(new)
        tc = root.find(f".//{qn('w:tbl')}/{qn('w:tr')}/{qn('w:tc')}")
        assert tc.find(f".//{qn('w:drawing')}") is not None
        # Layout not collapsed: the image paragraph is still its own <w:p>.
        assert len(tc.findall(qn("w:p"))) >= 3
        # The new text landed in the first text paragraph.
        cells = [e for e in docx_outline(new) if e.get("table") == 0]
        assert any("새 텍스트" in e["text"] for e in cells)

    def test_paragraph_replace_preserves_drawing_and_first_run_format(self):
        """Paragraph = [bold, plain, image, italic] runs. Replace keeps the
        drawing and the first run's bold formatting."""
        doc = Document()
        p = doc.add_paragraph()
        r_bold = p.add_run("볼드")
        r_bold.bold = True
        p.add_run("플레인")
        p.add_run().add_picture(io.BytesIO(_png_bytes()), width=Inches(0.5))
        r_italic = p.add_run("이탤릭")
        r_italic.italic = True
        base = _doc_bytes(doc)

        new, results = apply_docx_edits(
            base, [DocxEdit(action="replace", para=0, new_text="교체됨")]
        )
        assert results[0].status == "applied"

        assert _media_parts(new) == _media_parts(base)
        root = _document_xml(new)
        p_el = root.find(f"{qn('w:body')}/{qn('w:p')}")
        assert p_el.find(f".//{qn('w:drawing')}") is not None
        # Carrier run keeps its bold rPr and now holds the new text.
        first_run = p_el.find(qn("w:r"))
        assert first_run.find(f"{qn('w:rPr')}/{qn('w:b')}") is not None
        assert "".join(first_run.itertext()) == "교체됨"

    def test_paragraph_replace_keeps_nonempty_hyperlink(self):
        """A hyperlink-only paragraph: replace routes the new text INTO the
        hyperlink's run, so the w:hyperlink element survives (the old
        engine deleted hyperlinks wholesale)."""
        from docx.oxml import OxmlElement

        doc = Document()
        p = doc.add_paragraph()
        hyperlink = OxmlElement("w:hyperlink")
        run = OxmlElement("w:r")
        text_el = OxmlElement("w:t")
        text_el.text = "옛날링크"
        run.append(text_el)
        hyperlink.append(run)
        p._p.append(hyperlink)
        base = _doc_bytes(doc)

        new, results = apply_docx_edits(
            base, [DocxEdit(action="replace", para=0, new_text="새 링크 텍스트")]
        )
        assert results[0].status == "applied"

        root = _document_xml(new)
        p_el = root.find(f"{qn('w:body')}/{qn('w:p')}")
        hyperlink_el = p_el.find(qn("w:hyperlink"))
        assert hyperlink_el is not None, "hyperlink deleted wholesale"
        assert "새 링크 텍스트" in "".join(hyperlink_el.itertext())
        assert [e["text"] for e in docx_outline(new)] == ["새 링크 텍스트"]

    def test_paragraph_replace_touches_only_document_xml(self):
        """Byte-preservation: a single paragraph replace rewrites ONLY
        word/document.xml."""
        doc = Document()
        doc.add_paragraph("첫 문단")
        doc.add_paragraph("둘째 문단")
        base = _doc_bytes(doc)

        new, results = apply_docx_edits(
            base,
            [DocxEdit(action="replace", para=0, new_text="바뀐 첫 문단", old_text="첫 문단")],
        )
        assert results[0].status == "applied"
        assert _diff_parts(base, new) == {"word/document.xml"}

    def test_dead_hyperlink_stripped_after_replace(self):
        """A hyperlink NOT carrying the new text (paragraph has a direct
        text run) is emptied then stripped — matching the engine's
        no-stale-link intent."""
        from docx.oxml import OxmlElement

        doc = Document()
        p = doc.add_paragraph()
        p.add_run("본문 ")
        hyperlink = OxmlElement("w:hyperlink")
        run = OxmlElement("w:r")
        text_el = OxmlElement("w:t")
        text_el.text = "링크"
        run.append(text_el)
        hyperlink.append(run)
        p._p.append(hyperlink)
        base = _doc_bytes(doc)

        new, results = apply_docx_edits(
            base, [DocxEdit(action="replace", para=0, new_text="교체")]
        )
        assert results[0].status == "applied"
        root = _document_xml(new)
        p_el = root.find(f"{qn('w:body')}/{qn('w:p')}")
        assert p_el.find(qn("w:hyperlink")) is None  # dead link stripped
        assert [e["text"] for e in docx_outline(new)] == ["교체"]
