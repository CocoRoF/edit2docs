"""Deterministic DOCX building blocks (no LLM).

Markdown is the interchange format: the generator LLM writes markdown and
``docx_from_markdown`` renders it; ``docx_to_markdown`` goes the other way
(via mammoth) so editors can read a document back. Paragraph-level
addressing (``docx_outline`` / ``apply_docx_edits``) keeps untouched
paragraphs byte-identical — same philosophy as the PPTX recompose path.

Supported markdown subset (generation):
  # .. ###### headings · paragraphs · - / * bullets · 1. numbered lists ·
  **bold** / *italic* / `code` inline · pipe tables · > blockquote ·
  --- horizontal rule · ``` fenced code blocks (monospace paragraphs)
"""

from __future__ import annotations

import io
import re
from collections.abc import Iterable
from dataclasses import dataclass

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

__all__ = [
    "docx_from_markdown",
    "base_font_for_lang",
    "docx_to_markdown",
    "docx_to_html",
    "docx_preview",
    "docx_outline",
    "apply_docx_edits",
    "DocxEdit",
    "DocxEditResult",
]


# ---------------------------------------------------------------------------
# Markdown -> DOCX
# ---------------------------------------------------------------------------

_INLINE = re.compile(
    r"(\*\*(?P<bold>.+?)\*\*)|(\*(?P<italic>[^*]+?)\*)|(`(?P<code>[^`]+?)`)"
)


def _add_runs(paragraph, text: str) -> None:
    """Render **bold** / *italic* / `code` inline markup into runs."""
    pos = 0
    for match in _INLINE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        if match.group("bold") is not None:
            paragraph.add_run(match.group("bold")).bold = True
        elif match.group("italic") is not None:
            paragraph.add_run(match.group("italic")).italic = True
        else:
            run = paragraph.add_run(match.group("code"))
            run.font.name = "D2Coding"
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _is_table_row(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2


def _split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _is_separator_row(line: str) -> bool:
    """True only for genuine markdown alignment rows (---, :---:, ...).

    A permissive character-class match also swallowed real data rows like
    ``| - | - |`` — require every cell to be ``:?-{2,}:?``.
    """
    cells = _split_table_row(line)
    return bool(cells) and all(
        re.fullmatch(r":?-{2,}:?", cell) for cell in cells
    )


# Per-locale body typefaces for freshly generated documents. English-first
# default (Calibri, Word's own default face); CJK locales get their platform
# staples. Fragments rendered for insert_after inherit the TARGET document's
# Normal style on insertion, so this only shapes brand-new documents.
_BASE_FONTS = {
    "ko": "맑은 고딕",
    "ja": "Yu Gothic",
    "zh": "Microsoft YaHei",
}
_DEFAULT_BASE_FONT = "Calibri"


def base_font_for_lang(lang: str | None) -> str:
    """Default body typeface for a BCP-47 locale (en -> Calibri, ko -> 맑은 고딕)."""
    prefix = (lang or "").split("-")[0].lower()
    return _BASE_FONTS.get(prefix, _DEFAULT_BASE_FONT)


def docx_from_markdown(
    markdown: str, *, base_font: str | None = None, lang: str | None = None
) -> bytes:
    """Render the supported markdown subset into a .docx package.

    ``base_font`` wins when given; otherwise the face follows ``lang``
    (English-first: Calibri unless a CJK locale asks for its native face).
    """
    base_font = base_font or base_font_for_lang(lang)
    document = Document()
    style = document.styles["Normal"]
    style.font.name = base_font
    style.font.size = Pt(10.5)
    # East-Asian typeface must be set on the rPr explicitly.
    style.element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
        base_font,
    )

    lines = markdown.replace("\r\n", "\n").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Fenced code block -> monospace paragraphs.
        if stripped.startswith("```"):
            i += 1
            code_lines: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # closing fence
            for code_line in code_lines or [""]:
                p = document.add_paragraph()
                run = p.add_run(code_line)
                run.font.name = "D2Coding"
                run.font.size = Pt(9)
            continue

        # Pipe table.
        if _is_table_row(stripped):
            rows: list[list[str]] = []
            while i < len(lines) and _is_table_row(lines[i].strip()):
                if not _is_separator_row(lines[i]):
                    rows.append(_split_table_row(lines[i]))
                i += 1
            if rows:
                cols = max(len(r) for r in rows)
                table = document.add_table(rows=len(rows), cols=cols)
                table.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx in range(cols):
                        cell_text = row[c_idx] if c_idx < len(row) else ""
                        cell_p = table.rows[r_idx].cells[c_idx].paragraphs[0]
                        _add_runs(cell_p, cell_text)
                        if r_idx == 0:
                            for run in cell_p.runs:
                                run.bold = True
            continue

        # Heading.
        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            level = len(heading.group(1))
            p = document.add_heading("", level=min(level, 6))
            _add_runs(p, heading.group(2))
            i += 1
            continue

        # Horizontal rule -> centered separator.
        if re.match(r"^(-{3,}|\*{3,}|_{3,})$", stripped):
            p = document.add_paragraph("⸻")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Blockquote.
        if stripped.startswith(">"):
            p = document.add_paragraph(style="Intense Quote")
            _add_runs(p, stripped.lstrip("> ").strip())
            i += 1
            continue

        # Bullet / numbered list item.
        bullet = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet:
            p = document.add_paragraph(style="List Bullet")
            _add_runs(p, bullet.group(1))
            i += 1
            continue
        numbered = re.match(r"^\d+[.)]\s+(.*)$", stripped)
        if numbered:
            p = document.add_paragraph(style="List Number")
            _add_runs(p, numbered.group(1))
            i += 1
            continue

        # Plain paragraph (consume soft-wrapped continuation lines).
        chunk = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (
                not nxt
                or nxt.startswith(("#", ">", "```", "-", "*"))
                or _is_table_row(nxt)
                or re.match(r"^\d+[.)]\s", nxt)
            ):
                break
            chunk.append(nxt)
            i += 1
        p = document.add_paragraph()
        _add_runs(p, " ".join(chunk))

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# DOCX -> markdown / outline
# ---------------------------------------------------------------------------


def docx_to_markdown(content: bytes) -> str:
    """Convert a .docx to markdown via mammoth (structure-preserving)."""
    import mammoth
    from markdownify import markdownify

    html = mammoth.convert_to_html(io.BytesIO(content)).value
    return markdownify(html, heading_style="ATX").strip()


_HEADING_STYLE = re.compile(r"[Hh]eading\s*(\d)")


def docx_outline(content: bytes) -> list[dict]:
    """Paragraph outline with the addresses ``apply_docx_edits`` accepts.

    Entries: ``{"para": i, "style", "text"}`` for body paragraphs and
    ``{"table": t, "row": r, "col": c, "text"}`` for table cells — in
    **document order** (tables appear where they sit between paragraphs),
    so the edit planner sees the real structure of the page.

    Merged cells are reported once, at their top-left grid address (the
    same cell object python-docx returns for every covered position — and
    the address a ``replace`` edit actually mutates).

    Additive planner-visibility fields:

    * cells containing an inline image carry ``"has_image": True`` (and
      image-only cells are listed even with empty text, so the planner
      can see them);
    * charts referenced from the document part are appended as
      ``{"chart": i, "kind", "title"}`` entries (read-only, via
      contextifier's ChartModel).
    """
    from docx.oxml.ns import qn

    document = Document(io.BytesIO(content))
    paragraphs = document.paragraphs
    tables = document.tables
    outline: list[dict] = []
    para_index = -1
    table_index = -1
    for element in document.element.body:
        if element.tag == qn("w:p"):
            para_index += 1
            paragraph = paragraphs[para_index]
            text = paragraph.text.strip()
            if not text:
                continue
            style = paragraph.style.name if paragraph.style is not None else "Normal"
            outline.append({"para": para_index, "style": style, "text": text})
        elif element.tag == qn("w:tbl"):
            table_index += 1
            table = tables[table_index]
            # Merged cells share one w:tc element; report each once, at its
            # top-left grid address. Materialize all cells FIRST — holding
            # the wrappers keeps their lxml proxies alive, which is what
            # makes id()-identity of `_tc` reliable (a collected proxy's id
            # can be reused by a different element).
            grid = [
                (r, c, cell)
                for r, row in enumerate(table.rows)
                for c, cell in enumerate(row.cells)
            ]
            seen: set[int] = set()
            for r, c, cell in grid:
                key = id(cell._tc)
                if key in seen:
                    continue
                seen.add(key)
                text = cell.text.strip()
                has_image = (
                    next(cell._tc.iter(qn("w:drawing"), qn("w:pict")), None)
                    is not None
                )
                if text or has_image:
                    entry = {"table": table_index, "row": r, "col": c, "text": text}
                    if has_image:
                        entry["has_image"] = True
                    outline.append(entry)
    outline.extend(_chart_outline(content))
    return outline


def _chart_outline(content: bytes) -> list[dict]:
    """Read-only chart summaries via contextifier (best-effort: outline
    must never fail because a chart part is exotic)."""
    try:
        from contextifier import open_raw

        raw = open_raw(content, extension="docx")
        return [
            {"chart": i, "kind": chart.kind, "title": chart.title}
            for i, chart in enumerate(raw.charts)
        ]
    except Exception:
        return []


# ---------------------------------------------------------------------------
# Targeted edits
# ---------------------------------------------------------------------------


@dataclass
class DocxEdit:
    """One operation. ``action``:

    * ``replace`` — replace paragraph ``para``'s text (formatting of the
      first run survives) or table cell (``table``+``row``+``col``).
    * ``insert_after`` — insert markdown-rendered paragraphs after ``para``
      (para=-1 inserts at the start).
    * ``delete`` — remove paragraph ``para``.
    """

    action: str
    para: int | None = None
    table: int | None = None
    row: int | None = None
    col: int | None = None
    new_text: str = ""
    old_text: str | None = None
    markdown: str = ""


@dataclass
class DocxEditResult:
    action: str
    status: str  # applied | stale | not_found | invalid
    message: str = ""


def _normalize(text: str) -> str:
    return " ".join(text.split())


def apply_docx_edits(content: bytes, edits: Iterable[DocxEdit]) -> tuple[bytes, list[DocxEditResult]]:
    """Apply edits losslessly via contextifier's raw layer.

    Untouched parts stay byte-identical (only ``word/document.xml`` is
    rewritten). Replaces are run-preserving: the first text run keeps
    its formatting, in-paragraph/in-cell drawings and bookmarks survive,
    multi-paragraph cells keep their layout, and hyperlink elements are
    never deleted wholesale (dead ones emptied by the edit are stripped,
    matching the old engine's no-stale-link intent).

    Per-edit soft failures (like the PPTX text editor): ``old_text``
    guards replaces with a whitespace-normalized comparison.
    """
    from contextifier import open_raw

    try:
        raw = open_raw(content, extension="docx")
    except Exception as exc:
        raise ValueError(
            f"DOCX could not be opened: {exc}. DOCX 파일을 열 수 없습니다."
        ) from exc
    edit_list = list(edits)
    results: list[DocxEditResult | None] = [None] * len(edit_list)

    # Deletions/insertions shift indices — apply ops sorted by paragraph
    # DESCENDING so every edit's original address stays valid — but report
    # results in the CALLER'S order (clients correlate results[i] <-> edits[i]).
    ordered = sorted(
        enumerate(edit_list),
        key=lambda pair: (pair[1].para if pair[1].para is not None else 10**9),
        reverse=True,
    )
    for index, edit in ordered:
        results[index] = _apply_one(raw, edit)

    return raw.to_bytes(), [r for r in results if r is not None]


def _flatten(text: str) -> str:
    return " ".join(text.split("\n"))


def _apply_one(raw, edit: DocxEdit) -> DocxEditResult:
    if edit.action == "replace" and edit.table is not None:
        if edit.row is None or edit.col is None:
            return DocxEditResult(edit.action, "invalid", "table replace needs row+col")
        try:
            cell = raw.tables[edit.table].cell(edit.row, edit.col)
        except IndexError:
            return DocxEditResult(edit.action, "not_found", "table cell out of range")
        if edit.old_text is not None and _normalize(cell.text) != _normalize(edit.old_text):
            return DocxEditResult(edit.action, "stale", "cell text changed; refresh")
        # Run/layout-preserving: paragraphs holding drawings/pictures stay
        # untouched, other paragraphs are emptied but never removed.
        cell.set_text(_flatten(edit.new_text))
        return DocxEditResult(edit.action, "applied")

    paragraphs = raw.paragraphs
    para_ok = (
        edit.para is not None
        and 0 <= edit.para < len(paragraphs)
    ) or (edit.action == "insert_after" and edit.para == -1)
    if not para_ok:
        return DocxEditResult(edit.action, "not_found", "paragraph index out of range")

    if edit.action == "replace":
        current = paragraphs[edit.para].text
        if edit.old_text is not None and _normalize(current) != _normalize(edit.old_text):
            return DocxEditResult(edit.action, "stale", "paragraph text changed; refresh")
        raw.set_paragraph_text(edit.para, _flatten(edit.new_text))
        # Links whose text this edit emptied are dead — drop them (and
        # their now-unreferenced relationships); links that carry the new
        # text (hyperlink-only paragraphs) survive.
        raw.strip_empty_hyperlinks(edit.para)
        return DocxEditResult(edit.action, "applied")

    if edit.action == "delete":
        raw.delete_paragraph(edit.para)
        return DocxEditResult(edit.action, "applied")

    if edit.action == "insert_after":
        new_elements = _fragment_blocks(docx_from_markdown(edit.markdown or edit.new_text))
        if not new_elements:
            return DocxEditResult(edit.action, "invalid", "nothing to insert")
        document_part = raw.xml_part("word/document.xml")
        body = document_part.find("w:body")
        para_els = [p.element for p in paragraphs]
        if edit.para == -1:
            anchor = para_els[0] if para_els else None
            for element in new_elements:
                if anchor is None:
                    body.append(element)
                else:
                    anchor.addprevious(element)
        else:
            anchor = para_els[edit.para]
            for element in reversed(new_elements):
                anchor.addnext(element)
        document_part.mark_dirty()
        return DocxEditResult(edit.action, "applied")

    return DocxEditResult(edit.action, "invalid", f"unknown action {edit.action!r}")


def _fragment_blocks(fragment: bytes) -> list:
    """Body-level block elements (paragraphs AND tables) of a rendered
    markdown fragment, in document order; the section properties element
    and empty paragraphs are skipped (same rules as the python-docx-era
    grafting this replaces)."""
    import zipfile

    from contextifier.raw import qn
    from lxml import etree

    with zipfile.ZipFile(io.BytesIO(fragment)) as zf:
        root = etree.fromstring(zf.read("word/document.xml"))
    body = root.find(qn("w:body"))
    if body is None:
        return []
    blocks = []
    for element in body:
        if element.tag == qn("w:sectPr"):
            continue
        if element.tag == qn("w:p") and not "".join(element.itertext()).strip():
            continue
        blocks.append(element)
    return blocks


_SAFE_HREF = re.compile(r"^(https?:|mailto:)", re.IGNORECASE)
_ANCHOR_HREF = re.compile(r'<a\s+href="([^"]*)"')


def _sanitize_preview_html(html: str) -> str:
    """Neutralize unsafe URL schemes in anchors (javascript:, data:, ...).

    mammoth emits only structural tags, but hyperlink hrefs come verbatim
    from the docx relationships — a crafted document can carry
    ``javascript:`` URLs. Allowlist http/https/mailto; safe links open in
    a new tab so the studio SPA keeps its state.
    """

    def _fix(match: re.Match) -> str:
        href = match.group(1)
        if not _SAFE_HREF.match(href):
            return "<a"  # drop the href entirely
        return f'<a target="_blank" rel="noopener noreferrer" href="{href}"'

    return _ANCHOR_HREF.sub(_fix, html)


def docx_preview(content: bytes) -> tuple[str, list[dict]]:
    """Addressable display HTML + warnings for the studio preview.

    Primary path is the native renderer (:mod:`.docx_html`): document-order
    HTML where every body paragraph carries ``data-e2d-para`` and every
    table cell ``data-e2d-table`` / ``data-e2d-cell`` — the exact addresses
    ``docx_outline`` reports and the live-edit op stream targets, so the
    canvas can locate and highlight the region an edit touches (same
    convention as the PPTX preview's ``data-e2p-*`` tags).

    Any renderer failure falls back to the legacy mammoth conversion (with
    a warning entry) so a preview is always produced for a readable file.
    """
    from .docx_html import render_docx_html

    try:
        result = render_docx_html(content)
        return result.html, result.warnings
    except Exception as exc:
        html = _mammoth_html(content)
        return html, [
            {
                "code": "preview_native_render_failed",
                "message": (
                    f"Native DOCX renderer failed ({exc}); fell back to the "
                    "legacy converter. 기본 렌더러가 실패해 대체 변환기를 사용했습니다."
                ),
            }
        ]


def _mammoth_html(content: bytes) -> str:
    import mammoth

    return _sanitize_preview_html(mammoth.convert_to_html(io.BytesIO(content)).value)


def docx_to_html(content: bytes) -> str:
    """Convert a .docx to display HTML (addressable; see :func:`docx_preview`).

    Anchor hrefs are restricted to http/https/mailto in both the native
    renderer and the mammoth fallback, so the result is safe to inline in
    a styled preview container.
    """
    html, _warnings = docx_preview(content)
    return html
