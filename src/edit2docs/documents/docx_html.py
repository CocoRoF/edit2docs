"""Native, address-preserving DOCX -> display-HTML renderer (no LLM).

Why not mammoth: the studio needs the preview to be *addressable* — the
same paragraph / table-cell addresses that ``docx_outline`` reports and
``apply_docx_edits`` / the live-edit op stream consume. mammoth emits
clean semantic HTML but drops the mapping (and alignment, colors, merged
cells, empty paragraphs, page breaks), so the canvas could display the
document but never point at a region.

This renderer walks the document body **in document order** and stamps:

* ``data-e2d-para="i"``  — body paragraph *i* (python-docx
  ``document.paragraphs[i]``, the address ``replace``/``delete``/
  ``insert_after`` take),
* ``data-e2d-table="t"`` on ``<table>`` and ``data-e2d-cell="r,c"`` on
  cells — the ``table``+``row``+``col`` address of cell edits (merged
  cells carry their python-docx grid start, exactly what the editor
  targets),

mirroring the PPTX preview's ``data-e2p-shape`` / ``data-e2p-para``
convention (tools/apply_text_edits.py).

Fidelity carried per run: bold / italic / underline / strikethrough /
superscript / subscript / monospace (as ``<code>``) / highlight (as
``<mark>``) / font color; per paragraph: heading level (English and
Korean style names + ``w:outlineLvl``), list grouping (``<ul>``/``<ol>``
with ``w:ilvl`` nesting), blockquote styles, alignment; plus hyperlinks
(same allowlist policy as before), inline images (base64, size-capped),
page breaks, footnote references with a footnotes section, and empty
paragraphs (kept — they are real vertical rhythm in Word).

Any failure falls back to the legacy mammoth path with a warning entry,
so the preview never regresses below the old behavior.
"""

from __future__ import annotations

import base64
import io
import re
from dataclasses import dataclass, field
from html import escape

from docx import Document

__all__ = ["render_docx_html", "DocxHtmlResult"]

_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
_R = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"

_SAFE_HREF = re.compile(r"^(https?:|mailto:)", re.IGNORECASE)

# Base64 payload budget for embedded images per document. Beyond it,
# further images render as placeholders (with a warning) so a photo-heavy
# report cannot balloon the preview response.
_IMAGE_BUDGET_BYTES = 8 * 1024 * 1024

_HEADING_STYLE = re.compile(r"^(?:heading|제목)\s*(\d)$", re.IGNORECASE)
_MONO_FONTS = {
    "consolas", "courier new", "courier", "d2coding", "menlo", "monaco",
    "source code pro", "jetbrains mono", "nanumgothiccoding", "나눔고딕코딩",
}
_QUOTE_STYLES = {"quote", "intense quote", "인용", "강한 인용"}
_ORDERED_LIST_STYLE = re.compile(r"list\s*number|번호\s*매기기", re.IGNORECASE)
_LIST_STYLE = re.compile(r"list\s*(bullet|number|paragraph)|목록|글머리", re.IGNORECASE)

_ALIGN_CSS = {"center": "center", "right": "right", "both": "justify", "distribute": "justify"}

# w:highlight named colors -> CSS (the OOXML palette).
_HIGHLIGHT_CSS = {
    "yellow": "#ffff00", "green": "#00ff00", "cyan": "#00ffff", "magenta": "#ff00ff",
    "blue": "#0000ff", "red": "#ff0000", "darkBlue": "#00008b", "darkCyan": "#008b8b",
    "darkGreen": "#006400", "darkMagenta": "#8b008b", "darkRed": "#8b0000",
    "darkYellow": "#808000", "darkGray": "#a9a9a9", "lightGray": "#d3d3d3",
}


@dataclass
class DocxHtmlResult:
    html: str
    warnings: list[dict] = field(default_factory=list)


def render_docx_html(content: bytes) -> DocxHtmlResult:
    """Render a .docx to addressable display HTML.

    Raises ``ValueError`` when the bytes are not a readable DOCX package
    (bilingual message) — callers decide whether to fall back.
    """
    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:
        raise ValueError(
            f"DOCX could not be opened: {exc}. DOCX 파일을 열 수 없습니다."
        ) from exc
    renderer = _Renderer(document)
    return DocxHtmlResult(html=renderer.render(), warnings=renderer.warnings)


class _Renderer:
    def __init__(self, document) -> None:
        self.document = document
        self.warnings: list[dict] = []
        self._image_bytes_used = 0
        self._image_budget_warned = False
        self._footnote_refs: list[str] = []  # ids in first-reference order
        self._footnote_texts = self._load_footnotes()

    # -- top level ---------------------------------------------------------

    def render(self) -> str:
        out: list[str] = []
        list_stack: list[str] = []  # open list tags ("ul"/"ol"), one per level

        def close_lists(depth: int = 0) -> None:
            while len(list_stack) > depth:
                out.append(f"</{list_stack.pop()}>")

        para_index = -1
        table_index = -1
        for element in self.document.element.body:
            if element.tag == f"{_W}p":
                para_index += 1
                paragraph = _para_at(self.document, para_index)
                info = self._list_info(paragraph)
                if info is None:
                    close_lists()
                    out.append(self._render_paragraph(paragraph, para_index))
                    continue
                tag, level = info
                # Reconcile the open-list stack with (tag, level).
                while len(list_stack) > level + 1:
                    out.append(f"</{list_stack.pop()}>")
                while len(list_stack) < level + 1:
                    opening = tag if len(list_stack) == level else "ul"
                    list_stack.append(opening)
                    out.append(f"<{opening}>")
                if list_stack[level] != tag:
                    out.append(f"</{list_stack.pop()}>")
                    list_stack.append(tag)
                    out.append(f"<{tag}>")
                out.append(
                    f'<li data-e2d-para="{para_index}">'
                    f"{self._render_runs(paragraph._p)}</li>"
                )
            elif element.tag == f"{_W}tbl":
                close_lists()
                table_index += 1
                out.append(self._render_table(element, table_index))
            elif element.tag == f"{_W}sdt":
                # Block-level content control (TOC, cover page, ...). Its
                # paragraphs are not python-docx-addressable, so render the
                # content read-only — losing it entirely (what skipping
                # would do) is worse than losing its addresses.
                close_lists()
                out.append(self._render_sdt_block(element))
            elif element.tag == f"{_W}sectPr":
                continue
        close_lists()
        out.append(self._render_footnotes_section())
        return "".join(part for part in out if part)

    # -- paragraphs --------------------------------------------------------

    def _render_paragraph(self, paragraph, para_index: int) -> str:
        addr = f' data-e2d-para="{para_index}"'
        inner = self._render_runs(paragraph._p)
        style_name = (paragraph.style.name if paragraph.style is not None else "") or ""

        if not inner.strip() and not paragraph.text.strip():
            # Real empty paragraphs are Word's vertical rhythm — keep them
            # (mammoth dropped them, collapsing the page).
            return f'<p class="e2d-empty"{addr}><br/></p>'

        heading = _HEADING_STYLE.match(style_name.strip())
        level = None
        if heading:
            level = int(heading.group(1))
        else:
            outline_lvl = paragraph._p.find(f"{_W}pPr/{_W}outlineLvl")
            if outline_lvl is not None:
                try:
                    level = int(outline_lvl.get(f"{_W}val")) + 1
                except (TypeError, ValueError):
                    level = None
        if level is not None and 1 <= level <= 6:
            return f"<h{level}{addr}>{inner}</h{level}>"

        css = self._paragraph_css(paragraph)
        style_attr = f' style="{css}"' if css else ""
        if style_name.strip().lower() in _QUOTE_STYLES:
            return f"<blockquote{addr}{style_attr}><p>{inner}</p></blockquote>"
        return f"<p{addr}{style_attr}>{inner}</p>"

    def _paragraph_css(self, paragraph) -> str:
        try:
            alignment = paragraph.alignment  # None when inherited
        except ValueError:
            alignment = None
        if alignment is None:
            jc = paragraph._p.find(f"{_W}pPr/{_W}jc")
            value = jc.get(f"{_W}val") if jc is not None else None
        else:
            value = str(alignment).split(" ")[0].split(".")[-1].lower()
        css = _ALIGN_CSS.get((value or "").lower())
        return f"text-align:{css}" if css else ""

    def _list_info(self, paragraph) -> tuple[str, int] | None:
        """(``"ul"``/``"ol"``, 0-based level) when the paragraph is a list item."""
        style_name = (paragraph.style.name if paragraph.style is not None else "") or ""
        num_pr = paragraph._p.find(f"{_W}pPr/{_W}numPr")
        if num_pr is None and not _LIST_STYLE.search(style_name):
            return None
        level = 0
        if num_pr is not None:
            ilvl = num_pr.find(f"{_W}ilvl")
            if ilvl is not None:
                try:
                    level = max(0, min(int(ilvl.get(f"{_W}val")), 8))
                except (TypeError, ValueError):
                    level = 0
        tag = "ol" if self._is_ordered(paragraph, num_pr, style_name) else "ul"
        return tag, level

    def _is_ordered(self, paragraph, num_pr, style_name: str) -> bool:
        if _ORDERED_LIST_STYLE.search(style_name):
            return True
        if num_pr is None:
            return False
        try:
            num_id_el = num_pr.find(f"{_W}numId")
            ilvl_el = num_pr.find(f"{_W}ilvl")
            num_id = num_id_el.get(f"{_W}val") if num_id_el is not None else None
            ilvl = ilvl_el.get(f"{_W}val") if ilvl_el is not None else "0"
            fmt = self._num_fmt(num_id, ilvl)
            return fmt not in (None, "bullet", "none")
        except Exception:
            return False

    def _num_fmt(self, num_id: str | None, ilvl: str) -> str | None:
        """numFmt for (numId, ilvl) from the numbering part, if resolvable."""
        if num_id is None:
            return None
        numbering = getattr(self.document.part, "numbering_part", None)
        if numbering is None:
            return None
        root = numbering.element
        num = root.find(f'{_W}num[@{_W}numId="{num_id}"]')
        if num is None:
            return None
        abstract_ref = num.find(f"{_W}abstractNumId")
        if abstract_ref is None:
            return None
        abstract = root.find(
            f'{_W}abstractNum[@{_W}abstractNumId="{abstract_ref.get(f"{_W}val")}"]'
        )
        if abstract is None:
            return None
        lvl = abstract.find(f'{_W}lvl[@{_W}ilvl="{ilvl}"]')
        if lvl is None:
            return None
        fmt = lvl.find(f"{_W}numFmt")
        return fmt.get(f"{_W}val") if fmt is not None else None

    # -- runs & inline content ----------------------------------------------

    def _render_runs(self, p_element, *, part=None) -> str:
        """Render a ``w:p``'s inline content: runs, hyperlinks, images, breaks."""
        part = part or self.document.part
        out: list[str] = []
        for child in p_element:
            if child.tag == f"{_W}r":
                out.append(self._render_run(child, part))
            elif child.tag == f"{_W}hyperlink":
                out.append(self._render_hyperlink(child, part))
            elif child.tag in (
                f"{_W}ins",
                f"{_W}smartTag",
                f"{_W}sdt",
                f"{_W}sdtContent",
            ):
                # Tracked insertions / smart tags / content controls wrap
                # regular runs — render what's inside, ignore the wrapper.
                # (w:del — tracked deletions — is intentionally NOT rendered.)
                out.append(self._render_runs(child, part=part))
        return "".join(out)

    def _render_run(self, r_element, part) -> str:
        pieces: list[str] = []
        for child in r_element:
            tag = child.tag
            if tag == f"{_W}t":
                pieces.append(escape(child.text or ""))
            elif tag == f"{_W}tab":
                pieces.append("&emsp;")
            elif tag == f"{_W}br":
                if child.get(f"{_W}type") == "page":
                    pieces.append('<hr class="e2d-page-break"/>')
                else:
                    pieces.append("<br/>")
            elif tag == f"{_W}drawing" or tag == f"{_W}pict":
                pieces.append(self._render_image(child, part))
            elif tag == f"{_W}footnoteReference":
                pieces.append(self._render_footnote_ref(child))
        text = "".join(pieces)
        if not text:
            return ""
        return self._wrap_run_formatting(text, r_element.find(f"{_W}rPr"))

    def _wrap_run_formatting(self, text: str, rpr) -> str:
        if rpr is None:
            return text

        def on(tag: str) -> bool:
            el = rpr.find(f"{_W}{tag}")
            return el is not None and el.get(f"{_W}val") not in ("0", "false", "none")

        fonts = rpr.find(f"{_W}rFonts")
        if fonts is not None:
            names = {
                (fonts.get(f"{_W}{attr}") or "").lower()
                for attr in ("ascii", "hAnsi", "eastAsia", "cs")
            }
            if names & _MONO_FONTS:
                text = f"<code>{text}</code>"

        styles: list[str] = []
        color = rpr.find(f"{_W}color")
        if color is not None:
            value = color.get(f"{_W}val") or ""
            if re.fullmatch(r"[0-9A-Fa-f]{6}", value) and value.upper() != "000000":
                styles.append(f"color:#{value}")
        highlight = rpr.find(f"{_W}highlight")
        if highlight is not None:
            css = _HIGHLIGHT_CSS.get(highlight.get(f"{_W}val") or "")
            if css:
                text = f'<mark style="background-color:{css}">{text}</mark>'

        vert = rpr.find(f"{_W}vertAlign")
        if vert is not None and "e2d-footnote-ref" not in text:
            # (footnote markers are already superscripted by their own tag)
            value = vert.get(f"{_W}val")
            if value == "superscript":
                text = f"<sup>{text}</sup>"
            elif value == "subscript":
                text = f"<sub>{text}</sub>"

        if on("strike") or on("dstrike"):
            text = f"<s>{text}</s>"
        underline = rpr.find(f"{_W}u")
        if underline is not None and underline.get(f"{_W}val") not in ("none", None):
            text = f"<u>{text}</u>"
        if on("i") or on("iCs"):
            text = f"<em>{text}</em>"
        if on("b") or on("bCs"):
            text = f"<strong>{text}</strong>"
        if styles:
            text = f'<span style="{";".join(styles)}">{text}</span>'
        return text

    def _render_hyperlink(self, link_element, part) -> str:
        inner = "".join(
            self._render_run(child, part)
            for child in link_element
            if child.tag == f"{_W}r"
        )
        r_id = link_element.get(f"{_R}id")
        href = None
        if r_id:
            rel = part.rels.get(r_id)
            if rel is not None and rel.is_external:
                href = rel.target_ref
        if href and _SAFE_HREF.match(href):
            return (
                f'<a target="_blank" rel="noopener noreferrer" '
                f'href="{escape(href, quote=True)}">{inner}</a>'
            )
        # Internal anchors and unsafe schemes: keep the text, drop the link.
        return inner

    # -- images --------------------------------------------------------------

    def _render_image(self, drawing_element, part) -> str:
        blip = drawing_element.find(f".//{_A}blip")
        r_id = blip.get(f"{_R}embed") if blip is not None else None
        if not r_id:
            return ""
        try:
            image_part = part.related_parts[r_id]
            blob = image_part.blob
            content_type = image_part.content_type
        except (KeyError, AttributeError):
            return '<span class="e2d-image-missing">[이미지]</span>'
        if self._image_bytes_used + len(blob) > _IMAGE_BUDGET_BYTES:
            if not self._image_budget_warned:
                self._image_budget_warned = True
                self.warnings.append(
                    {
                        "code": "preview_images_truncated",
                        "message": (
                            "Embedded images exceed the preview budget; further "
                            "images are shown as placeholders. 이미지 용량 초과로 "
                            "일부 이미지는 자리표시자로 표시됩니다."
                        ),
                    }
                )
            return '<span class="e2d-image-omitted">[이미지 생략]</span>'
        self._image_bytes_used += len(blob)
        encoded = base64.b64encode(blob).decode("ascii")
        return (
            f'<img class="e2d-image" src="data:{content_type};base64,{encoded}" '
            f'alt="document image"/>'
        )

    # -- tables ----------------------------------------------------------------

    def _render_table(self, tbl_element, table_index: int, *, nested: bool = False) -> str:
        attrs = (
            ' class="e2d-table" data-e2d-nested="1"'
            if nested
            else f' class="e2d-table" data-e2d-table="{table_index}"'
        )
        rows_html: list[str] = []
        # rowspan bookkeeping: grid column -> remaining covered rows.
        vmerge_open: dict[int, dict] = {}
        rows = [child for child in tbl_element if child.tag == f"{_W}tr"]
        for r, tr in enumerate(rows):
            cells_html: list[str] = []
            grid_col = 0
            header_row = tr.find(f"{_W}trPr/{_W}tblHeader") is not None
            for tc in tr:
                if tc.tag != f"{_W}tc":
                    continue
                tc_pr = tc.find(f"{_W}tcPr")
                span_el = tc_pr.find(f"{_W}gridSpan") if tc_pr is not None else None
                try:
                    col_span = max(1, int(span_el.get(f"{_W}val"))) if span_el is not None else 1
                except (TypeError, ValueError):
                    col_span = 1
                vmerge = tc_pr.find(f"{_W}vMerge") if tc_pr is not None else None
                if vmerge is not None and vmerge.get(f"{_W}val") in (None, "continue"):
                    # Covered by the cell above: count it into that cell's
                    # rowspan and emit nothing.
                    open_cell = vmerge_open.get(grid_col)
                    if open_cell is not None:
                        open_cell["rowspan"] += 1
                    grid_col += col_span
                    continue

                cell_inner = self._render_cell_blocks(tc)
                tag = "th" if header_row else "td"
                cell = {
                    "tag": tag,
                    "row": r,
                    "col": grid_col,
                    "colspan": col_span,
                    "rowspan": 1,
                    "html": cell_inner,
                    "nested": nested,
                }
                if vmerge is not None:  # restart: may grow via later rows
                    vmerge_open[grid_col] = cell
                cells_html.append(cell)
                grid_col += col_span
            rows_html.append(cells_html)

        body: list[str] = [f"<table{attrs}>"]
        for cells in rows_html:
            body.append("<tr>")
            for cell in cells:
                span_attrs = ""
                if cell["colspan"] > 1:
                    span_attrs += f' colspan="{cell["colspan"]}"'
                if cell["rowspan"] > 1:
                    span_attrs += f' rowspan="{cell["rowspan"]}"'
                addr = (
                    ""
                    if cell["nested"]
                    else f' data-e2d-cell="{cell["row"]},{cell["col"]}"'
                )
                body.append(
                    f"<{cell['tag']}{addr}{span_attrs}>{cell['html']}</{cell['tag']}>"
                )
            body.append("</tr>")
        body.append("</table>")
        return "".join(body)

    def _render_cell_blocks(self, tc_element) -> str:
        parts: list[str] = []
        for child in tc_element:
            if child.tag == f"{_W}p":
                inner = self._render_runs(child)
                parts.append(f"<p>{inner}</p>" if inner else "")
            elif child.tag == f"{_W}tbl":
                parts.append(self._render_table(child, -1, nested=True))
        blocks = [p for p in parts if p]
        if not blocks:
            return "<p></p>"
        return "".join(blocks)

    def _render_sdt_block(self, sdt_element) -> str:
        content = sdt_element.find(f"{_W}sdtContent")
        if content is None:
            return ""
        parts: list[str] = ['<div class="e2d-sdt">']
        for child in content:
            if child.tag == f"{_W}p":
                inner = self._render_runs(child)
                parts.append(f"<p>{inner}</p>" if inner else "")
            elif child.tag == f"{_W}tbl":
                parts.append(self._render_table(child, -1, nested=True))
            elif child.tag == f"{_W}sdt":
                parts.append(self._render_sdt_block(child))
        parts.append("</div>")
        rendered = [p for p in parts if p]
        return "".join(rendered) if len(rendered) > 2 else ""

    # -- footnotes ----------------------------------------------------------

    def _load_footnotes(self) -> dict[str, str]:
        """id -> text from the footnotes part (empty when absent).

        python-docx has no footnotes API and loads the part generically,
        so parse the raw XML from the part blob.
        """
        try:
            from lxml import etree

            for rel in self.document.part.rels.values():
                if rel.reltype.endswith("/footnotes"):
                    root = etree.fromstring(rel.target_part.blob)
                    notes: dict[str, str] = {}
                    for note in root.findall(f"{_W}footnote"):
                        note_id = note.get(f"{_W}id")
                        if note_id in ("-1", "0", None):  # separators
                            continue
                        text = "".join(note.itertext()).strip()
                        if text:
                            notes[note_id] = text
                    return notes
        except Exception:
            pass
        return {}

    def _render_footnote_ref(self, ref_element) -> str:
        note_id = ref_element.get(f"{_W}id")
        if note_id is None or note_id not in self._footnote_texts:
            return ""
        if note_id not in self._footnote_refs:
            self._footnote_refs.append(note_id)
        number = self._footnote_refs.index(note_id) + 1
        return f'<sup class="e2d-footnote-ref">[{number}]</sup>'

    def _render_footnotes_section(self) -> str:
        if not self._footnote_refs:
            return ""
        items = "".join(
            f"<li>{escape(self._footnote_texts[note_id])}</li>"
            for note_id in self._footnote_refs
        )
        return f'<section class="e2d-footnotes"><ol>{items}</ol></section>'


def _para_at(document, index: int):
    """document.paragraphs[index] without re-building the list every call."""
    cached = getattr(document, "_e2d_para_cache", None)
    if cached is None:
        cached = document.paragraphs
        document._e2d_para_cache = cached
    return cached[index]
